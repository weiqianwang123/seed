import docx,random
from tkinter import *



list0=[["逢考必过","考神护体"]]

def listcreate(a):  #传入一个文件操作后产生的列表a，返回一个二维列表，每个单元是一组问答
    num = len(a.paragraphs)
    list_x=[]
    for i in range(0,num-1,2):
        list_y=[]
        if a.paragraphs[i].text!='':
            list_y.append(a.paragraphs[i].text)
            list_y.append(a.paragraphs[i+1].text)
            list_x.append(list_y)

    return list_x
class workinglist:
    def __init__(self,list0,list1,list2,list3,list4,list5,list6,list7,list8,root):
        self.list1=list1
        self.list2=list2
        self.list3=list3
        self.list4=list4
        self.list5=list5
        self.list6=list6
        self.list7=list7
        self.list8=list8
        self.root=root
        self.work_list = list0
        self.i=0
    def new_wind(self,i):     #此函数用来产生一个新窗口，显示问题的答案,i是问题的序号
        NewWind=Toplevel(self.root)
        NewWind.geometry('1000x600')
        NewWind.title('答案')
        NewLabel=Message(NewWind,text=self.work_list[i][1],width=500,font=("黑体",35))
        NewLabel.place(relx=0.2,rely=0.2)



    def list_change1(self):     #修改当前所用工作列表为需要列表，属于菜单命令所需要函数
         self.work_list = self.list1
         self.i = 0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change2(self):
         self.work_list = self.list2
         self.i = 0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change3(self):
         self.work_list = self.list3
         self.i=0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change4(self):
         self.work_list = self.list4
         self.i = 0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change5(self):
         self.work_list = self.list5
         self.i=0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change6(self):
         self.work_list = self.list6
         self.i=0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))
    def list_change7(self):
         self.work_list = self.list7
         self.i=0
         QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))

    def list_change8(self):
        self.work_list = self.list8
        self.i = 0
        QuestionLabel.config(text=self.work_list[self.i][0], font=("黑体", 40))

    def update(self,idx):  # 定时器函数
        frame = self.frames[idx]
        idx += 1  # 下一帧的序号：在0,1,2,3,4,5之间循环(共6帧)
        self.photolabel.configure(image=frame)  # 显示当前帧的图片
        root.after(100, self.update, idx % 40)


    def func1(self):           #随机顺序按钮，随机产生一个序号的问题，并修改label
        self.i=random.randint(0,len(self.work_list)-1)
        QuestionLabel.config(text=self.work_list[self.i][0],font=("黑体",40))

    def func2(self):            #答案按钮，打开新窗口并显示问题的答案
        self.new_wind(self.i)

    def func3(self):        #顺序按钮，按顺序显示下一个问题
        if self.i<len(self.work_list)-1:
            self.i=self.i+1
            QuestionLabel.config(text=self.work_list[self.i][0],font=("黑体",40))
        else:
            QuestionLabel.config(text='已经到底啦')






    def power(self):
        NewWind = Toplevel(self.root)
        NewWind.geometry('400x200')
        NewWind.title('玄学押题')
        self.work_list = self.list8
        self.i = 0
        endbt=Button(NewWind,text="一发即中",command=self.func1)
        endbt.pack()
        numIdx = 40
        self.frames = [PhotoImage(file='R-C.gif', format='gif -index %i' % (i)) for i in range(numIdx)]
        self.photolabel = Label(NewWind, fg='black', font=("黑体", 80), width=1000, image=self.frames[0])
        self.photolabel.pack()
        NewWind.after(0, self.update, 0)











#读取word文档，并且转换成可用列表,并且以这些列表为基本元素，引用workinglist类


list1=listcreate(docx.Document("0.docx"))
list2=listcreate(docx.Document("1.docx"))
list3=listcreate(docx.Document("2.docx"))
list4=listcreate(docx.Document("3.docx"))
list5=listcreate(docx.Document("4.docx"))
list6=listcreate(docx.Document("5.docx"))
list7=listcreate(docx.Document("6.docx"))
list8=listcreate(docx.Document("7.docx"))







#先进行窗口的初始化
root = Tk()
root.geometry('500x500')
root.title("seed")
a=workinglist(list0,list1,list2,list3,list4,list5,list6,list7,list8,root)
QuestionLabel=Message(root,text='prod by WQW',fg='black',font=("黑体",80),width=1000)
QuestionLabel.pack()



#再添加页面元素按钮
bt1=Button(root,text="随机",command=a.func1)
bt1.place(relx=0.05, rely=0.7, relwidth=0.3, relheight=0.1)
bt2=Button(root,text="答案",command=a.func2)
bt2.place(relx=0.4, rely=0.7, relwidth=0.3, relheight=0.1)
bt3=Button(root,text="顺序",command=a.func3)
bt3.place(relx=0.75, rely=0.7, relwidth=0.3, relheight=0.1)


#添加菜单组件
def popmenu(event):
    mainmenu.post(event.x_root, event.u_root)
mainmenu = Menu(root)
menuFile = Menu(mainmenu)
mainmenu.add_cascade(label="单词表",menu=menuFile)
menuFile.add_command(label="词库",command=a.list_change1)
menuFile.add_command(label="收藏夹",command=a.list_change2)
#menuFile.add_command(label="军事思想",command=a.list_change3)
#menuFile.add_command(label="现代战争",command=a.list_change4)
#menuFile.add_command(label="信息化装备",command=a.list_change5)
#menuFile.add_command(label="第五单元",command=a.list_change6)
#menuFile.add_command(label="第六单元",command=a.list_change7)
#menuFile.add_command(label="全部题库",command=a.list_change8)
menuFile.add_command(label="退出",command=root.destroy)
#choice=Menu(mainmenu)
#mainmenu.add_cascade(label="玄学押题模式",menu=choice)
#choice.add_command(label="临考之前，点击此键",command=a.power)


root.config(menu=mainmenu)
root.bind('Button-3',popmenu) # 根窗体绑定鼠标右击响应事件


root.mainloop()